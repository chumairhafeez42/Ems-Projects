import mimetypes
from io import FileIO, BytesIO

from docx import Document


class FileExtractor:
    """
    Wrapper for extracting file contents to string
    """

    def __init__(
            self,
            file_obj=None,
            filename: str or None = None,
    ):
        """
        The extractor wrapper will initialize by assinging the filename to the
        object's file property; if a file-like object is provided instead of a
        name, then a file_ext arg will be required.
        """
        if filename:
            self.file = FileIO(filename)
            self.file_ext = filename.split(".")[-1]
        else:
            if hasattr(file_obj, "name"):
                self.file = file_obj
                self.file_ext = file_obj.name.split(".")[-1]
            else:
                raise Exception(
                    "Your file object does not contain a name attribute. Please"
                    " add a name attribute with a file extension, and try "
                    "again. Need the file ext. data for mime-typing."
                )

    @staticmethod
    def get_file_type(file):
        mime_type = mimetypes.guess_type(file)[0]
        return mime_type.split("/")[1]

    def get_contents(self):
        with self.file as f:
            f.seek(0, 0)
            return f.read()

    def PdfFileRead(self):
        """
        This current code provides a workaround in case MuPDF (a dependency
        for PyMuPDF) is not usable in the development environment. For such
        instances, the module relies on PyPDF2 to extract text data. However,
        because of the likelihood of white spaces being rampant in the
        extracted string data, those characters get filtered out.
        """

        contents = self.get_contents()

        try:
            import fitz

            pdf_file = fitz.Document(stream=contents, filetype="pdf")
            raw_text = [ele.get_text("text") for ele in pdf_file]
            text = "".join(raw_text)
        # else:
        except Exception:
            import PyPDF2

            pdf_reader = PyPDF2.PdfFileReader(contents)
            raw_text = [ele.extractText() for ele in pdf_reader.pages]
            text = "".join(raw_text)
        return text

    def DocxFileRead(self):
        contents = self.get_contents()
        f_stream = BytesIO(contents)
        document = Document(f_stream)
        raw_text = [p.text for p in document.paragraphs]
        text = "\n".join(raw_text)
        return text

    def TextFileRead(self):
        if isinstance(self.file, FileIO): return self.file.read().decode('utf-8')
        with open(self.file) as f:
            return f.read()


class MyWordLoader(object):
    def __init__(self, file_obj=None, filename: str or None = None):
        self.file = FileExtractor(file_obj, filename)

    def file_load(self):
        file_type = self.file.file_ext
        # file_type = file_loc.split('.')[-1]

        # file_types_tup = ('pdf', 'docx', 'doc', 'txt', 'text')
        file_types_tup = ("pdf", "docx", "txt", "text")
        if file_type in file_types_tup:
            if file_type == file_types_tup[0]:
                text = self.file.PdfFileRead()
            elif file_type == file_types_tup[1]:
                text = self.file.DocxFileRead()
            # elif file_type == file_types_tup[2]:
            #     text = DocFileRead(self.text)
            else:
                text = self.file.TextFileRead()
            return text
        else:
            mime_type = self.file.get_file_type(self.file.name)
            print(
                f"You are using an incorrect file format for file submissions.\n\
            Please upload a .docx/.doc/.txt/.pdf file OR!\n\
            Note the mimetype of your submitted data and submit an error \
            report to github with the following: {mime_type}"
            )


def TextSpitter(file_obj=None, filename: str or None = None):
    return MyWordLoader(file_obj=file_obj, filename=filename).file_load()
